from docx import Document
from docx.shared import Pt, RGBColor
import io
import os
import re
import tempfile
from docx2pdf import convert
from datetime import datetime
import pythoncom 

class ExportService:

    # ---------- 1. DOCUMENT SETTINGS ----------
    @staticmethod
    def get_document_meta(doc_type):
        dt = doc_type.lower().strip() if doc_type else ""
        if 'quiz' in dt:
            return {"marks": "10", "duration": "10 Minutes", "header": "Q1. Choose the best answers for the following MCQs. [10]"}
        elif 'assignment' in dt:
            return {"marks": "10", "duration": "N/A", "header": "Q1. Answer the following questions in detail. [10]"}
        elif 'mid' in dt:
            return {"marks": "20", "duration": "1 Hour", "header": "Q1. Attempt all questions. Section A (Objective) & Section B (Subjective)."}
        elif 'final' in dt:
            return {"marks": "40", "duration": "2 Hours", "header": "Q1. Attempt all questions. Section A (Objective) & Section B (Subjective)."}
        else:
            return {"marks": "10", "duration": "10 Mins", "header": "Questions:"}

    @staticmethod
    def get_department_name(program):
        p = str(program).upper()
        if "AI" in p or "ARTIFICIAL" in p:
            return "DEPARTMENT OF ARTIFICIAL INTELLIGENCE"
        elif "SE" in p or "SOFTWARE" in p:
            return "DEPARTMENT OF SOFTWARE ENGINEERING"
        elif "DS" in p or "DATA" in p:
            return "DEPARTMENT OF DATA SCIENCE"
        return "DEPARTMENT OF COMPUTER SCIENCE"

    # ---------- 2. TEXT CLEANER (REGEX POWERED) ----------
    @staticmethod
    def clean_text(text, include_answers=True):
        """Cleans AI output and optionally removes answers using Regex."""
        if not text: return ""

        # 1. Global Removal of Bold/Italic/Code markers
        # Removing **, __, ##, and the specific instruction parentheticals
        text = re.sub(r'\*\*|__', '', text) 
        text = re.sub(r'\(.*don.*t use topic.*\)', '', text, flags=re.IGNORECASE)

        lines = text.split('\n')
        filtered_lines = []
        
        for line in lines:
            l = line.strip()
            if not l: continue 
            
            # 2. Aggressively strip leading symbols (Bullet points, Hash headers, Dashes)
            # e.g., "# Question 1" becomes "Question 1"
            # e.g., "* A)" becomes "A)"
            l = re.sub(r'^[\#\*\-\s]+', '', l)

            # 3. Skip Garbage Lines (Separators or Labels)
            if re.match(r'^[-=_]{3,}$', l): continue # Lines like "---" or "==="
            if re.match(r'^(Title|Subject|Topic|Task|Date):', l, re.IGNORECASE): continue

            # 4. Skip Conversational Filler
            lower = l.lower()
            if lower.startswith("here is") or lower.startswith("sure, here") or \
               lower.startswith("as your professor") or "i have crafted" in lower:
                continue

            # 5. HIDE ANSWERS (Robust Regex Check)
            # Matches: "Correct Answer:", "Answer:", "Ans:", "Correct Option:"
            if not include_answers:
                if re.search(r'^(correct\s*)?(answer|option|ans)\s*[:\-]', lower):
                    continue
            
            filtered_lines.append(l)
        
        # Rejoin and return
        return "\n".join(filtered_lines)

    # ---------- 3. TEXT REPLACER ----------
    @staticmethod
    def replace_text_in_paragraph(paragraph, replacements):
        if not paragraph.text: return

        # 1. Try Simple Run Replacement
        for run in paragraph.runs:
            for key, val in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, str(val))
                    run.bold = True
                    run.font.name = 'Times New Roman'
        
        # 2. Fallback Replacement
        text = paragraph.text
        updated = False
        for key, val in replacements.items():
            if key in text:
                text = text.replace(key, str(val))
                updated = True
        
        if updated:
            paragraph.text = text
            paragraph.style.font.name = 'Times New Roman'
            paragraph.style.font.size = Pt(11)
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                for val in replacements.values():
                    if str(val) in run.text:
                        run.bold = True

    # ---------- 4. MAIN GENERATOR ----------
    @staticmethod
    def create_university_doc(content_text, course_name, course_code, teacher_name, doc_type, program_name="BS-CS", include_answers=True):
        
        template_path = "Template.docx"
        if not os.path.exists(template_path):
            doc = Document()
            doc.add_paragraph("Error: Template.docx not found.")
            stream = io.BytesIO()
            doc.save(stream)
            stream.seek(0)
            return stream

        doc = Document(template_path)
        meta = ExportService.get_document_meta(doc_type)
        
        replacements = {
            "{{DEPARTMENT}}": ExportService.get_department_name(program_name),
            "{{INSTRUCTOR}}": teacher_name,
            "{{SESSION}}": f"{doc_type.capitalize()} (Fall-{datetime.now().year})",
            "{{SUBJECT}}": course_name,
            "{{CODE}}": course_code,
            "{{MARKS}}": meta["marks"],
            "{{PROGRAM}}": program_name,
            "{{SHIFT}}": "Morning",
            "{{SECTION}}": "A",
            "{{DATE}}": datetime.now().strftime("%d %B %Y"),
            "{{DURATION}}": meta["duration"]
        }

        # 1. Fill Headers
        for p in doc.paragraphs:
            ExportService.replace_text_in_paragraph(p, replacements)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        ExportService.replace_text_in_paragraph(p, replacements)

        # 2. Fill Content
        content_filled = False
        question_counter = 1 

        for i, p in enumerate(doc.paragraphs):
            if "{{CONTENT}}" in p.text:
                p.text = "" 
                content_filled = True
                p.paragraph_format.space_after = Pt(0)

                run = p.add_run(meta["header"]) 
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

                if content_text:
                    cleaned_content = ExportService.clean_text(content_text, include_answers)
                    
                    for line in cleaned_content.split("\n"):
                        line = line.strip()
                        if not line: continue
                        
                        # --- A. DETECT ANSWERS (Backup Check) ---
                        # Even if clean_text missed it, we double check here before printing
                        is_answer = re.search(r'^(correct\s*)?(answer|option|ans)\s*[:\-]', line, re.IGNORECASE)
                        if is_answer:
                            if include_answers:
                                ans_p = doc.add_paragraph()
                                ans_p.paragraph_format.space_before = Pt(0)
                                ans_p.paragraph_format.space_after = Pt(6)
                                run = ans_p.add_run(line)
                                run.bold = True
                                run.font.color.rgb = RGBColor(0, 0, 0)
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(11)
                            continue # Skip this line if it's an answer (printed or hidden)

                        # --- B. DETECT QUESTIONS ---
                        # Matches "1.", "Q1", "Question 1"
                        match = re.match(r'^(?:Question|Q)?\s*(\d+)[\.\):]\s*(.*)', line, re.IGNORECASE)
                        if match:
                            q_text = match.group(2).strip()
                            
                            # Gap between questions
                            if question_counter > 1:
                                empty_p = doc.add_paragraph()
                                empty_p.paragraph_format.space_after = Pt(12) 

                            # Question Label
                            q_label = doc.add_paragraph()
                            q_label.paragraph_format.space_after = Pt(0)
                            run = q_label.add_run(f"Question {question_counter}:")
                            run.bold = True
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)

                            # Question Text
                            if q_text:
                                q_body = doc.add_paragraph()
                                q_body.paragraph_format.space_after = Pt(6)
                                run = q_body.add_run(q_text)
                                run.bold = True 
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(11)
                            
                            question_counter += 1
                        
                        # --- C. DETECT SECTIONS ---
                        elif line.startswith("SECTION"):
                            doc.add_paragraph()
                            new_p = doc.add_paragraph()
                            run = new_p.add_run(line)
                            run.bold = True
                            run.font.name = 'Times New Roman'
                        
                        # --- D. DETECT OPTIONS ---
                        else:
                            # Standardize option printing
                            p_opt = doc.add_paragraph()
                            p_opt.paragraph_format.space_after = Pt(0)
                            run = p_opt.add_run(line)
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
                break 

        if not content_filled:
            doc.add_paragraph(meta["header"])
            doc.add_paragraph(content_text if content_text else "No Content Generated")

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        return stream

    @staticmethod
    def docx_stream_to_pdf(docx_stream):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
            tmp_docx.write(docx_stream.getvalue())
            tmp_docx_path = tmp_docx.name
        
        tmp_pdf_path = tmp_docx_path.replace(".docx", ".pdf")
        try:
            pythoncom.CoInitialize()
            convert(tmp_docx_path, tmp_pdf_path)
            with open(tmp_pdf_path, "rb") as f:
                pdf_data = f.read()
            return io.BytesIO(pdf_data)
        except Exception as e:
            print(f"PDF Conversion Error: {e}")
            return None
        finally:
            if os.path.exists(tmp_docx_path): os.remove(tmp_docx_path)
            if os.path.exists(tmp_pdf_path): os.remove(tmp_pdf_path)
            pythoncom.CoUninitialize()